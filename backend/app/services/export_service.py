import json
import logging
import os
from typing import Literal

from app.core.config import settings

logger = logging.getLogger(__name__)

EXPORT_DIR = os.path.join(settings.UPLOAD_DIR, "exports")


class ExportService:
    def __init__(self):
        os.makedirs(EXPORT_DIR, exist_ok=True)
        logger.info("ExportService initialized")

    async def export(
        self,
        text: str,
        format: str,
        filename: str,
    ) -> str:
        os.makedirs(EXPORT_DIR, exist_ok=True)
        if format == "txt":
            return self._export_txt(text, filename)
        elif format == "pdf":
            return self._export_pdf(text, filename)
        elif format == "docx":
            return self._export_docx(text, filename)
        elif format == "json":
            return self._export_json(text, filename)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _export_txt(self, text: str, filename: str) -> str:
        path = os.path.join(EXPORT_DIR, f"{filename}.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)
        return path

    def _export_pdf(self, text: str, filename: str) -> str:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        path = os.path.join(EXPORT_DIR, f"{filename}.pdf")
        c = canvas.Canvas(path, pagesize=A4)
        width, height = A4
        c.setFont("Helvetica", 12)
        margin = 50
        y = height - margin
        for line in text.split("\n"):
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 12)
                y = height - margin
            c.drawString(margin, y, line)
            y -= 18
        c.save()
        return path

    def _export_docx(self, text: str, filename: str) -> str:
        from docx import Document

        path = os.path.join(EXPORT_DIR, f"{filename}.docx")
        doc = Document()
        doc.add_heading("Braille Conversion Result", 0)
        for para in text.split("\n"):
            doc.add_paragraph(para)
        doc.save(path)
        return path

    def _export_json(self, text: str, filename: str) -> str:
        path = os.path.join(EXPORT_DIR, f"{filename}.json")
        payload = {
            "recognized_text": text,
            "word_count": len(text.split()),
            "char_count": len(text),
        }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2)
        return path